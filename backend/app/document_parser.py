"""
Document Parser Module

Extracts text content from uploaded documents in multiple formats:
- TXT: Direct read (UTF-8)
- PDF: Text-based extraction via PyPDF2
- DOCX: Paragraph extraction via python-docx
"""

import io
import logging
from pathlib import Path
from typing import Tuple

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# PDF text extraction
# ---------------------------------------------------------------------------

def _extract_text_from_pdf(content: bytes) -> str:
    """
    Extract text from a PDF file using PyPDF2.
    Combines text from all pages into a single string.
    """
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise RuntimeError("PyPDF2 is required for PDF parsing. Install it with: pip install PyPDF2")

    reader = PdfReader(io.BytesIO(content))
    all_text: list[str] = []

    for page in reader.pages:
        page_text = (page.extract_text() or "").strip()
        if page_text:
            all_text.append(page_text)

    return "\n\n".join(all_text)


# ---------------------------------------------------------------------------
# DOCX text extraction
# ---------------------------------------------------------------------------

def _extract_text_from_docx(content: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx is required for DOCX parsing. Install it with: pip install python-docx")

    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


# ---------------------------------------------------------------------------
# TXT (passthrough)
# ---------------------------------------------------------------------------

def _extract_text_from_txt(content: bytes) -> str:
    """Decode a plain-text file (UTF-8 with fallback to latin-1)."""
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

PARSERS = {
    ".txt": _extract_text_from_txt,
    ".pdf": _extract_text_from_pdf,
    ".docx": _extract_text_from_docx,
}


def extract_text(filename: str, content: bytes) -> Tuple[str, str]:
    """
    Extract text from a document based on its file extension.

    Args:
        filename: Original filename (used to determine format).
        content: Raw file bytes.

    Returns:
        A tuple of (extracted_text, output_filename) where output_filename
        always has a .txt extension (for storage in the knowledge base).

    Raises:
        ValueError: If the file extension is not supported.
        RuntimeError: If a required library is missing.
    """
    ext = Path(filename).suffix.lower()

    parser = PARSERS.get(ext)
    if parser is None:
        supported = ", ".join(sorted(PARSERS.keys()))
        raise ValueError(f"Unsupported file type '{ext}'. Supported: {supported}")

    text = parser(content)

    if not text.strip():
        raise ValueError(
            f"No text could be extracted from '{filename}'. "
            "The file may be a scanned PDF with no selectable text."
        )

    # Output filename: replace extension with .txt
    stem = Path(filename).stem
    output_filename = f"{stem}.txt"

    return text, output_filename
